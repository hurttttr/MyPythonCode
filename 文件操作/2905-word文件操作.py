import docx
import random
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

doc = docx.Document()

doc.styles['Normal'].font.name = '华文行楷'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '华文行楷')

for i in range(3):
    p = doc.add_paragraph()
    for j in range(5):
        run = p.add_run('文本颜色测试(test)')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(5+i*5)

        run.font.bold = True
        run.font.italic = True

        color = (random.randint(0, 255) for i in range(3))
        run.font.color.rgb = RGBColor(*color)

doc.add_paragraph()
doc.save(r'文件操作\test.docx')
