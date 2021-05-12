from os import P_WAIT
import openpyxl
import random
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor

wb = openpyxl.load_workbook('工资统计.xlsx')
sheet = wb['工作量汇总']

lst = []
for index, values in enumerate(sheet.rows, 1):
    t = list(map(lambda x: x.value, values))
    lst.append(t)

doc = docx.Document()

p = doc.add_paragraph('工资表')
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p.runs[0].font.size = Pt(30)
p.runs[0].font.bold = True

table = doc.add_table(rows=13, cols=6, style='Table Grid')
for i in range(13):
    for j in range(6):
        value = lst[i][j]
        if type(value) == int:
            value = str(value)
        elif type(value) == float:
            value = '{:.2f}'.format(value)
        cell = table.cell(i, j)
        cell.text = value

        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        color = (random.randint(0, 255)for i in range(3))
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(*color)
doc.save('table.docx')
